from docx import Document

count = 0

variables = {
    "$EMPLOYEENAME$": ["Ayturan Fətəlisoy Famil oğlu", "Məmmədli Ərtürk Nəsimi oğlu"],
    "$EMPLOYEEWORKPLACE$": ["ƏƏSMN", "SOCAR"],
    "$EMPLOYEETITLE$": ["Specialist", "Engineer"],
    "$EMPLOYEEADDRESS$": ["Quzanlı, Ağdam", "Ş. Mirzəyev, Bakı"],
    "$EMPLOYEEPHONE$": ["+994-555555555", "+994-777777777"],
    "$EMPLOYEEEMAIL$": ["ayturan@example.com", "erturk@meselcun.com"],
    "$BIRTHDATE$": ["31 May 1993", "9 İyun 1992"],
    "$SALARYAMOUNT$": ["9,312 AZN", "8,712 AZN"]
}
    
while count < len(variables['$EMPLOYEENAME$']):
    
    template_file_path = 'C:/Users/HP/Documents/Python3/Test.docx'
    output_file_path = 'C:/Users/HP/Documents/Python3/%s.docx' % variables["$EMPLOYEENAME$"][count]

    template_document = Document(template_file_path)

    def replace_text_in_paragraph(paragraph, key, value):
        #print(paragraph.text)
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                #print('item -----', item.text)
                #print('key ------', key)
                if item.text in key:
                    #print('-----Cool-----')
                    item.text = item.text.replace(key, value)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value[count])

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value[count])

    template_document.save(output_file_path)
    
    count += 1