from docx import Document
import os

template_file_path = 'C:/Users/HP/Documents/Python3/Test.docx'
output_file_path = 'C:/Users/HP/Documents/Python3/Result.docx'

variables = {
    "$EMPLOYEENAME$": "Ayturan Fətəlisoy Famil oğlu",
    "$EMPLOYEEWORKPLACE$": "ƏƏSMN",
    "$EMPLOYEETITLE$": "Specialist",
    "$EMPLOYEEADDRESS$": "Quzanlı, Ağdam",
    "$EMPLOYEEPHONE$": "+994-505000000",
    "$EMPLOYEEEMAIL$": "ayturan@example.com",
    "$BIRTHDATE$": "31 May, 1993",
    "$SALARYAMOUNT$": "10,000 AZN"
}

template_document = Document(template_file_path)

def replace_text_in_paragraph(paragraph, key, value):
    #print(paragraph.text)
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            #print('item -----', item.text)
            #print('key ------', key)
            if item.text in key:
                #print('-----Cool-----')
                item.text = item.text.replace(key, value)

for variable_key, variable_value in variables.items():
    for paragraph in template_document.paragraphs:
        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    for table in template_document.tables:
        for col in table.columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

template_document.save(output_file_path)
